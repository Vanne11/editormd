# exporter.py - versión corregida con extensión automática

import os
from PySide6.QtWidgets import QFileDialog, QMessageBox
from PySide6.QtPrintSupport import QPrinter
from PySide6.QtGui import QTextDocument
import markdown

# Dependencias opcionales
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False


class Exporter:
    @staticmethod
    def export_dialog(parent, text):
        """Muestra un diálogo para exportar a distintos formatos"""

        # Convertir Markdown a HTML (para exportar con Qt PDF)
        html_content = markdown.markdown(text, extensions=['extra', 'codehilite', 'tables'])

        # Filtros de exportación según dependencias disponibles
        filters = [
            "Markdown (*.md)",
            "Texto (*.txt)"
        ]
        if DOCX_AVAILABLE:
            filters.append("DOCX (*.docx)")
        if PYPANDOC_AVAILABLE:
            filters.extend([
                "ODT (*.odt)",
                "PDF (Pandoc) (*.pdf)"
            ])
        filters.append("PDF (Qt - Recomendado) (*.pdf)")

        # Guardar archivo
        path, selected_filter = QFileDialog.getSaveFileName(
            parent,
            "Exportar archivo",
            "",
            ";;".join(filters)
        )

        if not path:
            return

        # 🔧 Forzar extensión si el usuario no la escribió
        ext_map = {
            "Markdown (*.md)": ".md",
            "Texto (*.txt)": ".txt",
            "DOCX (*.docx)": ".docx",
            "ODT (*.odt)": ".odt",
            "PDF (Pandoc) (*.pdf)": ".pdf",
            "PDF (Qt - Recomendado) (*.pdf)": ".pdf",
        }
        if selected_filter in ext_map and not path.lower().endswith(ext_map[selected_filter]):
            path += ext_map[selected_filter]

        try:
            if path.endswith(".docx") and DOCX_AVAILABLE:
                Exporter.to_docx(path, text)

            elif path.endswith(".odt") and PYPANDOC_AVAILABLE:
                Exporter.to_odt(path, text)

            elif path.endswith(".pdf") and "Pandoc" in selected_filter and PYPANDOC_AVAILABLE:
                Exporter.to_pdf_pandoc(path, text)

            elif path.endswith(".pdf"):
                # Siempre funciona, no requiere nada externo
                Exporter.to_pdf_qt(path, html_content)

            elif path.endswith(".md") or path.endswith(".txt"):
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)

            else:
                raise ValueError(
                    "⚠️ Falta una dependencia para este formato.\n\n"
                    "Ejemplos:\n"
                    " - .docx → requiere 'python-docx'\n"
                    " - .odt / .pdf (Pandoc) → requieren 'pypandoc' y 'pandoc'"
                )

            QMessageBox.information(parent, "✅ Éxito", f"Archivo exportado correctamente:\n{path}")

        except Exception as e:
            QMessageBox.critical(parent, "❌ Error", f"No se pudo exportar:\n\n{str(e)}")
            print(f"Error detallado: {e}")

    # ==== Exportadores individuales ====
    @staticmethod
    def to_docx(path, text):
        """Exporta a DOCX usando python-docx"""
        doc = Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(path)

    @staticmethod
    def to_odt(path, text):
        """Exporta a ODT usando pypandoc"""
        pypandoc.convert_text(text, "odt", format="md", outputfile=path, extra_args=["--standalone"])

    @staticmethod
    def to_pdf_pandoc(path, text):
        """Exporta a PDF usando pypandoc (requiere Pandoc + LaTeX)"""
        pypandoc.convert_text(text, "pdf", format="md", outputfile=path, extra_args=["--standalone"])

    @staticmethod
    def to_pdf_qt(path, html_content):
        """Exporta a PDF usando PySide6 (no requiere nada externo)"""
        printer = QPrinter()
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)

        document = QTextDocument()
        document.setHtml(html_content)
        document.print(printer)
